"""
Data Ingestion Script
=====================
Loads documents, performs semantic chunking, transforms to unified schema,
generates embeddings, and uploads to Qdrant.

Process:
1. Load .docx files with LangChain
2. Semantic chunking (pattern-based)
3. Parse Bug Reports (structured)
4. Parse User Feedbacks (infer metadata with LLM)
5. Transform to UnifiedDocument
6. Generate embeddings with nomic-embed-text
7. Upload to Qdrant
"""

import sys
import re
import json
import time
from pathlib import Path
from typing import List, Dict, Any
from datetime import datetime

# Add parent directory to path
sys.path.append(str(Path(__file__).parent.parent))

from docx import Document as DocxDocument
from langchain.schema import Document
from langchain_community.embeddings import OllamaEmbeddings
from langchain_community.vectorstores import Qdrant
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, VectorParams
import httpx

from app.config import settings
from app.models.schemas import UnifiedDocument


# ============================================================================
# CONFIGURATION
# ============================================================================

BUG_REPORT_PATH = Path("/app/data/raw/ai_test_bug_report.docx")
USER_FEEDBACK_PATH = Path("/app/data/raw/ai_test_user_feedback.docx")
PROCESSED_OUTPUT = Path("/app/data/processed/unified_documents.json")


# ============================================================================
# STEP 1: LOAD DOCUMENTS
# ============================================================================

def load_docx_content(filepath: Path) -> str:
    """Load content from .docx file using python-docx"""
    print(f" Loading: {filepath.name}")
    
    doc = DocxDocument(filepath)
    full_text = []
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            full_text.append(paragraph.text)
    
    content = "\n".join(full_text)
    print(f"   / Loaded {len(content)} characters")
    return content


# ============================================================================
# STEP 2: SEMANTIC CHUNKING
# ============================================================================

def semantic_chunking_bug_report(text: str) -> List[str]:
    """
    Semantic chunking for Bug Reports.
    Split by 'Bug #X' pattern to preserve complete bug reports.
    """
    print(" Semantic chunking: Bug Reports")
    
    # Split by "Bug #" pattern
    pattern = r"(Bug #\d+)"
    parts = re.split(pattern, text)
    
    chunks = []
    for i in range(1, len(parts), 2):  # Start from 1, step by 2
        if i < len(parts):
            bug_header = parts[i]  # "Bug #X"
            bug_content = parts[i + 1] if i + 1 < len(parts) else ""
            full_chunk = f"{bug_header}\n{bug_content.strip()}"
            chunks.append(full_chunk)
    
    print(f"   / Created {len(chunks)} chunks")
    return chunks


def semantic_chunking_user_feedback(text: str) -> List[str]:
    """
    Semantic chunking for User Feedbacks.
    Split by 'Feedback #X:' pattern to preserve complete feedbacks.
    """
    print("/ Semantic chunking: User Feedbacks")
    
    # Split by "Feedback #" pattern
    pattern = r"(Feedback #\d+:)"
    parts = re.split(pattern, text)
    
    chunks = []
    for i in range(1, len(parts), 2):
        if i < len(parts):
            feedback_header = parts[i]  # "Feedback #X:"
            feedback_content = parts[i + 1] if i + 1 < len(parts) else ""
            full_chunk = f"{feedback_header} {feedback_content.strip()}"
            chunks.append(full_chunk)
    
    print(f"   / Created {len(chunks)} chunks")
    return chunks


# ============================================================================
# STEP 3: PARSE BUG REPORTS (Structured)
# ============================================================================

def parse_bug_report_chunk(chunk: str) -> Dict[str, Any]:
    """Parse a single bug report chunk into structured fields"""
    
    lines = chunk.split("\n")
    bug_data = {
        "bug_id": "",
        "title": "",
        "description": "",
        "steps_to_reproduce": "",
        "environment": "",
        "severity": "",
        "proposed_fix": ""
    }
    
    # Extract Bug ID from first line
    bug_match = re.search(r"Bug #(\d+)", lines[0])
    if bug_match:
        bug_data["bug_id"] = bug_match.group(1)
    
    current_field = None
    current_content = []
    
    for line in lines[1:]:
        line = line.strip()
        if not line:
            continue
        
        # Detect field headers
        if line.startswith("Title:"):
            if current_field:
                bug_data[current_field] = " ".join(current_content).strip()
            current_field = "title"
            current_content = [line.replace("Title:", "").strip()]
        elif line.startswith("Description:"):
            if current_field:
                bug_data[current_field] = " ".join(current_content).strip()
            current_field = "description"
            current_content = [line.replace("Description:", "").strip()]
        elif line.startswith("Steps to Reproduce:"):
            if current_field:
                bug_data[current_field] = " ".join(current_content).strip()
            current_field = "steps_to_reproduce"
            current_content = [line.replace("Steps to Reproduce:", "").strip()]
        elif line.startswith("Environment:"):
            if current_field:
                bug_data[current_field] = " ".join(current_content).strip()
            current_field = "environment"
            current_content = [line.replace("Environment:", "").strip()]
        elif line.startswith("Severity:"):
            if current_field:
                bug_data[current_field] = " ".join(current_content).strip()
            current_field = "severity"
            current_content = [line.replace("Severity:", "").strip()]
        elif line.startswith("Proposed Fix:"):
            if current_field:
                bug_data[current_field] = " ".join(current_content).strip()
            current_field = "proposed_fix"
            current_content = [line.replace("Proposed Fix:", "").strip()]
        else:
            # Continue current field
            if current_field:
                current_content.append(line)
    
    # Don't forget the last field
    if current_field:
        bug_data[current_field] = " ".join(current_content).strip()
    
    return bug_data


def extract_category_from_title(title: str) -> str:
    """Extract category from bug title using keyword matching"""
    
    title_lower = title.lower()
    
    category_keywords = {
        "Document Management": ["upload", "download", "document", "file", "pdf"],
        "Search & Filtering": ["search", "filter", "query", "results", "pagination"],
        "UI/UX": ["button", "icon", "layout", "display", "color", "alignment", "tooltip", "overlap"],
        "Performance": ["slow", "loading", "speed", "performance", "lag"],
        "Mobile": ["mobile", "phone", "tablet", "touch", "responsive", "orientation"],
        "Authentication": ["login", "password", "reset", "authentication"],
        "Notifications": ["notification", "email", "alert"],
        "Collaboration": ["share", "sharing", "collaboration", "real-time"],
        "Backend": ["server", "database", "API", "backend", "integration"],
    }
    
    for category, keywords in category_keywords.items():
        if any(keyword in title_lower for keyword in keywords):
            return category
    
    return "General"


# ============================================================================
# STEP 4: PARSE USER FEEDBACKS (Unstructured + LLM Inference)
# ============================================================================

def parse_user_feedback_chunk(chunk: str) -> Dict[str, Any]:
    """Parse a single user feedback chunk"""
    
    # Extract Feedback ID
    feedback_match = re.search(r"Feedback #(\d+):", chunk)
    feedback_id = feedback_match.group(1) if feedback_match else "0"
    
    # Extract content (everything after "Feedback #X:")
    content = re.sub(r"Feedback #\d+:\s*", "", chunk).strip()
    
    return {
        "feedback_id": feedback_id,
        "content": content
    }


def infer_metadata_with_llm(content: str) -> Dict[str, str]:
    """
    Use LLM (llama3.2) to infer category and severity from user feedback.
    This is the KEY step that makes unstructured data structured!
    """
    
    prompt = f"""You are an expert at analyzing user feedback. Analyze the following feedback and extract:
1. Category: Choose ONE from [Document Management, Search & Filtering, UI/UX, Performance, Mobile, Authentication, Notifications, Collaboration, Backend, General]
2. Severity: Choose ONE from [Low, Medium, High]

User Feedback: "{content}"

Respond ONLY with valid JSON in this exact format:
{{"category": "...", "severity": "..."}}

JSON:"""

    try:
        response = httpx.post(
            f"{settings.OLLAMA_URL}/api/generate",
            json={
                "model": settings.LLM_MODEL,
                "prompt": prompt,
                "stream": False,
                "temperature": 0.3,  # Lower temperature for consistent classification
            },
            timeout=60.0  # Increased to 60s for cold start
        )
        response.raise_for_status()
        
        result = response.json()
        response_text = result["response"].strip()
        
        # Clean response (remove markdown code blocks if present)
        response_text = re.sub(r"```json\s*", "", response_text)
        response_text = re.sub(r"```\s*", "", response_text)
        response_text = response_text.strip()
        
        # Parse JSON
        inferred = json.loads(response_text)
        
        # Validate
        valid_categories = ["Document Management", "Search & Filtering", "UI/UX", 
                          "Performance", "Mobile", "Authentication", "Notifications", 
                          "Collaboration", "Backend", "General"]
        valid_severities = ["Low", "Medium", "High"]
        
        if inferred["category"] not in valid_categories:
            inferred["category"] = "General"
        if inferred["severity"] not in valid_severities:
            inferred["severity"] = "Medium"
        
        return inferred
        
    except Exception as e:
        print(f"   !  LLM inference failed: {e}, using defaults")
        return {"category": "General", "severity": "Medium"}


# ============================================================================
# STEP 5: TRANSFORM TO UNIFIED DOCUMENT
# ============================================================================

def transform_bug_reports(chunks: List[str]) -> List[UnifiedDocument]:
    """Transform bug report chunks to UnifiedDocuments"""
    
    print("\n Transforming Bug Reports to UnifiedDocuments")
    unified_docs = []
    
    for chunk in chunks:
        try:
            bug_data = parse_bug_report_chunk(chunk)
            
            # Skip if parsing failed
            if not bug_data["bug_id"]:
                continue
            
            category = extract_category_from_title(bug_data["title"])
            
            # Create embedding text
            embedding_text = f"""Category: {category}
Severity: {bug_data['severity']}
Title: {bug_data['title']}
Description: {bug_data['description']}
Environment: {bug_data['environment']}"""
            
            unified_doc = UnifiedDocument(
                id=f"BUG-{bug_data['bug_id'].zfill(3)}",
                source_type="bug_report",
                title=bug_data['title'],
                content=bug_data['description'],
                category=category,
                severity=bug_data['severity'],
                environment=bug_data['environment'],
                steps_to_reproduce=bug_data['steps_to_reproduce'],
                proposed_fix=bug_data['proposed_fix'],
                embedding_text=embedding_text
            )
            
            unified_docs.append(unified_doc)
            
        except Exception as e:
            print(f"   !  Failed to parse chunk: {e}")
            continue
    
    print(f"   / Transformed {len(unified_docs)} bug reports")
    return unified_docs


def transform_user_feedbacks(chunks: List[str]) -> List[UnifiedDocument]:
    """Transform user feedback chunks to UnifiedDocuments (with LLM inference)"""
    
    print("\n Transforming User Feedbacks to UnifiedDocuments")
    print("    Using LLM to infer metadata...")
    
    unified_docs = []
    
    for i, chunk in enumerate(chunks, 1):
        try:
            feedback_data = parse_user_feedback_chunk(chunk)
            
            # Infer metadata with LLM (KEY STEP!)
            print(f"   [{i}/{len(chunks)}] Inferring metadata...", end=" ")
            inferred = infer_metadata_with_llm(feedback_data['content'])
            print(f"Category: {inferred['category']}, Severity: {inferred['severity']}")
            
            # Create embedding text
            embedding_text = f"""Category: {inferred['category']}
Severity: {inferred['severity']}
Feedback: {feedback_data['content']}"""
            
            unified_doc = UnifiedDocument(
                id=f"FEEDBACK-{feedback_data['feedback_id'].zfill(3)}",
                source_type="user_feedback",
                title=None,  # User feedback doesn't have titles
                content=feedback_data['content'],
                category=inferred['category'],
                severity=inferred['severity'],
                environment=None,
                steps_to_reproduce=None,
                proposed_fix=None,
                embedding_text=embedding_text
            )
            
            unified_docs.append(unified_doc)
            
        except Exception as e:
            print(f"   !  Failed to process feedback: {e}")
            continue
    
    print(f"   / Transformed {len(unified_docs)} user feedbacks")
    return unified_docs


# ============================================================================
# STEP 6: GENERATE EMBEDDINGS & UPLOAD TO QDRANT
# ============================================================================

def upload_to_qdrant(unified_docs: List[UnifiedDocument]):
    """Generate embeddings and upload to Qdrant using LangChain"""
    
    print("\n Uploading to Qdrant")
    print(f"    Total documents: {len(unified_docs)}")
    
    # Initialize Qdrant client with increased timeout
    qdrant_client = QdrantClient(
        url=settings.QDRANT_URL,
        timeout=120  # 2 minutes timeout
    )
    
    # Delete collection if exists (fresh start)
    try:
        qdrant_client.delete_collection(collection_name=settings.QDRANT_COLLECTION)
        print(f"     Deleted existing collection: {settings.QDRANT_COLLECTION}")
        time.sleep(2)  # Wait for deletion to complete
    except Exception as e:
        print(f"     No existing collection to delete")
    
    # Create collection with retry logic
    max_retries = 3
    for attempt in range(max_retries):
        try:
            qdrant_client.create_collection(
                collection_name=settings.QDRANT_COLLECTION,
                vectors_config=VectorParams(size=settings.VECTOR_SIZE, distance=Distance.COSINE)
            )
            print(f"   / Created collection: {settings.QDRANT_COLLECTION}")
            break
        except Exception as e:
            if attempt < max_retries - 1:
                wait_time = (attempt + 1) * 5
                print(f"   !  Attempt {attempt + 1} failed, retrying in {wait_time}s...")
                time.sleep(wait_time)
            else:
                print(f"   ! Failed to create collection after {max_retries} attempts")
                raise
    
    # Initialize embeddings model
    print(f"    Initializing embeddings model: {settings.EMBEDDING_MODEL}")
    embeddings_model = OllamaEmbeddings(
        base_url=settings.OLLAMA_URL,
        model=settings.EMBEDDING_MODEL
    )
    
    # Convert to LangChain Documents
    langchain_docs = []
    for doc in unified_docs:
        langchain_doc = Document(
            page_content=doc.embedding_text,
            metadata={
                "id": doc.id,
                "source_type": doc.source_type,
                "category": doc.category,
                "severity": doc.severity,
                "title": doc.title or "",
                "content": doc.content,
                "environment": doc.environment or "",
                "steps_to_reproduce": doc.steps_to_reproduce or "",
                "proposed_fix": doc.proposed_fix or "",
            }
        )
        langchain_docs.append(langchain_doc)
    
    # Upload to Qdrant with LangChain
    print(f"    Generating embeddings and uploading...")
    start_time = time.time()
    
    vector_store = Qdrant.from_documents(
        documents=langchain_docs,
        embedding=embeddings_model,
        url=settings.QDRANT_URL,
        collection_name=settings.QDRANT_COLLECTION,
    )
    
    elapsed = time.time() - start_time
    print(f"   / Upload complete! Time: {elapsed:.2f}s")
    print(f"     {len(langchain_docs)} documents in Qdrant")


# ============================================================================
# STEP 7: SAVE PROCESSED DATA (Backup)
# ============================================================================

def save_processed_data(unified_docs: List[UnifiedDocument]):
    """Save processed data as JSON backup"""
    
    print("\n/ Saving processed data")
    
    # Convert to dict
    data = [doc.model_dump() for doc in unified_docs]
    
    # Ensure directory exists
    PROCESSED_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    
    # Save as JSON
    with open(PROCESSED_OUTPUT, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=2, ensure_ascii=False, default=str)
    
    print(f"   / Saved to: {PROCESSED_OUTPUT}")
    print(f"     {len(data)} documents")


# ============================================================================
# MAIN EXECUTION
# ============================================================================

def main():
    """Main ingestion pipeline"""
    
    print("=" * 70)
    print(" DATA INGESTION PIPELINE")
    print("=" * 70)
    
    start_time = time.time()
    
    # Step 1: Load Documents
    print("\n STEP 1: Loading Documents")
    bug_report_text = load_docx_content(BUG_REPORT_PATH)
    user_feedback_text = load_docx_content(USER_FEEDBACK_PATH)
    
    # Step 2: Semantic Chunking
    print("\n STEP 2: Semantic Chunking")
    bug_chunks = semantic_chunking_bug_report(bug_report_text)
    feedback_chunks = semantic_chunking_user_feedback(user_feedback_text)
    
    # Step 3: Transform to Unified Schema
    print("\n STEP 3: Transform to Unified Schema")
    bug_docs = transform_bug_reports(bug_chunks)
    feedback_docs = transform_user_feedbacks(feedback_chunks)
    
    all_unified_docs = bug_docs + feedback_docs
    print(f"\n    Total unified documents: {len(all_unified_docs)}")
    
    # Step 4: Upload to Qdrant
    print("\n STEP 4: Upload to Qdrant")
    upload_to_qdrant(all_unified_docs)
    
    # Step 5: Save Backup
    print("\n STEP 5: Save Backup")
    save_processed_data(all_unified_docs)
    
    # Summary
    elapsed = time.time() - start_time
    print("\n" + "=" * 70)
    print("/ INGESTION COMPLETE!")
    print("=" * 70)
    print(f"  Statistics:")
    print(f"   - Bug Reports: {len(bug_docs)}")
    print(f"   - User Feedbacks: {len(feedback_docs)}")
    print(f"   - Total Documents: {len(all_unified_docs)}")
    print(f"   - Collection: {settings.QDRANT_COLLECTION}")
    print(f"   - Vector Size: {settings.VECTOR_SIZE}")
    print(f"   - Total Time: {elapsed:.2f}s")
    print("=" * 70)


if __name__ == "__main__":
    main()